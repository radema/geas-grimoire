#!/usr/bin/env python3
"""
Markdown to Word Document Converter

Converts a markdown file with YAML frontmatter into a styled Word document
using a template for look & feel.

Dependencies: python-docx, pyyaml, mistune

Usage:
    python md_to_docx.py --input content.md --template template.docx --output report.docx
"""

import argparse
import re
import sys
from pathlib import Path

import yaml

try:
    import mistune
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError as e:
    print(f"Missing dependency: {e.name}")
    print("Install with: pip install python-docx pyyaml mistune")
    sys.exit(1)


def parse_frontmatter(content: str) -> tuple[dict, str]:
    """Extract YAML frontmatter and markdown body from content."""
    frontmatter = {}
    body = content
    
    if content.startswith("---"):
        parts = content.split("---", 2)
        if len(parts) >= 3:
            try:
                frontmatter = yaml.safe_load(parts[1]) or {}
            except yaml.YAMLError as e:
                print(f"Warning: Failed to parse YAML frontmatter: {e}")
            body = parts[2].strip()
    
    return frontmatter, body


def replace_placeholders(text: str, placeholders: dict) -> str:
    """Replace {{placeholder}} tokens with values from dict."""
    if not placeholders:
        return text
    
    def replacer(match):
        key = match.group(1).strip()
        if key in placeholders:
            return str(placeholders[key])
        print(f"Warning: Placeholder '{{{{{key}}}}}' not found in YAML")
        return match.group(0)  # Leave unchanged
    
    return re.sub(r"\{\{(.+?)\}\}", replacer, text)


class DocxRenderer(mistune.HTMLRenderer):
    """Custom renderer that builds Word document elements."""
    
    HTML_TAG_REGEX = re.compile(r"(<strong>|</strong>|<em>|</em>|<code>|</code>)")

    def __init__(self, doc: Document, placeholders: dict, base_path: Path):
        super().__init__()
        self.doc = doc
        self.placeholders = placeholders
        self.base_path = base_path
        self._current_list_type = None
        self._list_stack = []
    
    def text(self, text: str) -> str:
        return replace_placeholders(text, self.placeholders)
    
    def paragraph(self, text: str) -> str:
        # Skip empty paragraphs
        clean_text = re.sub(r"<[^>]+>", "", text).strip()
        if clean_text:
            p = self.doc.add_paragraph()
            self._add_formatted_text(p, text)
        return ""
    
    def heading(self, text: str, level: int, **attrs) -> str:
        clean_text = re.sub(r"<[^>]+>", "", text).strip()
        clean_text = replace_placeholders(clean_text, self.placeholders)
        
        # Map to Word heading styles
        style_name = f"Heading {min(level, 9)}"
        try:
            self.doc.add_heading(clean_text, level=min(level, 9))
        except Exception:
            # Fallback if heading style doesn't exist
            p = self.doc.add_paragraph(clean_text)
            p.style = "Normal"
        return ""
    
    def strong(self, text: str) -> str:
        return f"<strong>{text}</strong>"
    
    def emphasis(self, text: str) -> str:
        return f"<em>{text}</em>"
    
    def codespan(self, text: str) -> str:
        return f"<code>{text}</code>"
    
    def list(self, body: str, ordered: bool, **attrs) -> str:
        return body
    
    def list_item(self, text: str, **attrs) -> str:
        clean_text = re.sub(r"<[^>]+>", "", text).strip()
        clean_text = replace_placeholders(clean_text, self.placeholders)
        
        # Add as bullet list item
        p = self.doc.add_paragraph(clean_text, style="List Bullet")
        return ""
    
    def image(self, alt: str, url: str, title: str = None) -> str:
        """Insert image into document."""
        image_path = self.base_path / url
        
        if not image_path.exists():
            # Try relative to current dir
            image_path = Path(url)
        
        if image_path.exists():
            try:
                self.doc.add_picture(str(image_path), width=Inches(5.5))
                # Center the image
                last_para = self.doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add caption if alt text provided
                if alt and alt != "Chart":
                    caption = self.doc.add_paragraph(alt)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.runs[0].italic = True
            except Exception as e:
                print(f"Warning: Could not insert image {url}: {e}")
                self.doc.add_paragraph(f"[Image not found: {url}]")
        else:
            print(f"Warning: Image not found: {image_path}")
            self.doc.add_paragraph(f"[Image not found: {url}]")
        
        return ""
    
    def _add_formatted_text(self, paragraph, html_text: str):
        """Parse HTML-like formatting and add runs to paragraph."""
        # Simple parser for <strong>, <em>, <code> tags
        parts = self.HTML_TAG_REGEX.split(html_text)
        
        bold = False
        italic = False
        code = False
        
        for part in parts:
            if part == "<strong>":
                bold = True
            elif part == "</strong>":
                bold = False
            elif part == "<em>":
                italic = True
            elif part == "</em>":
                italic = False
            elif part == "<code>":
                code = True
            elif part == "</code>":
                code = False
            elif part.strip():
                text = replace_placeholders(part, self.placeholders)
                run = paragraph.add_run(text)
                run.bold = bold
                run.italic = italic
                if code:
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)


def parse_markdown_table(lines: list[str], doc: Document, placeholders: dict):
    """Parse markdown table lines and add to document."""
    if len(lines) < 2:
        return
    
    # Parse header
    header_cells = [c.strip() for c in lines[0].strip("|").split("|")]
    num_cols = len(header_cells)
    
    # Skip separator line (line 1)
    data_lines = lines[2:]
    rows = []
    for line in data_lines:
        if "|" in line:
            cells = [c.strip() for c in line.strip("|").split("|")]
            rows.append(cells)
    
    # Create table
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = "Table Grid"
    
    # Fill header
    header_row = table.rows[0]
    for i, cell_text in enumerate(header_cells):
        if i < len(header_row.cells):
            cell = header_row.cells[i]
            cell.text = replace_placeholders(cell_text, placeholders)
            # Bold header
            for run in cell.paragraphs[0].runs:
                run.bold = True
    
    # Fill data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < len(row.cells):
                row.cells[col_idx].text = replace_placeholders(cell_text, placeholders)


def convert_markdown_to_docx(
    input_path: Path,
    template_path: Path,
    output_path: Path
) -> bool:
    """Main conversion function."""
    
    # Read input
    try:
        content = input_path.read_text(encoding="utf-8")
    except Exception as e:
        print(f"Error reading input file: {e}")
        return False
    
    # Parse frontmatter
    frontmatter, body = parse_frontmatter(content)
    placeholders = frontmatter.get("placeholders", {})
    
    # Load template or create new document
    if template_path.exists():
        try:
            doc = Document(str(template_path))
        except Exception as e:
            print(f"Error loading template: {e}")
            print("Creating new document instead.")
            doc = Document()
    else:
        print(f"Warning: Template '{template_path}' not found. Creating blank document.")
        doc = Document()
    
    # Process markdown in sections to handle tables specially
    lines = body.split("\n")
    current_block = []
    in_table = False
    table_lines = []
    
    # Initialize renderer and parser once
    renderer = DocxRenderer(doc, placeholders, input_path.parent)
    md = mistune.create_markdown(renderer=renderer)

    for line in lines:
        # Detect table start
        if "|" in line and not in_table:
            # Process any pending markdown
            if current_block:
                block_text = "\n".join(current_block)
                md(block_text)
                current_block = []
            in_table = True
            table_lines = [line]
        elif in_table:
            if "|" in line or (line.strip().startswith("|") or line.strip().startswith("-")):
                table_lines.append(line)
            else:
                # End of table
                parse_markdown_table(table_lines, doc, placeholders)
                in_table = False
                table_lines = []
                current_block = [line]
        else:
            current_block.append(line)
    
    # Process remaining content
    if in_table and table_lines:
        parse_markdown_table(table_lines, doc, placeholders)
    elif current_block:
        block_text = "\n".join(current_block)
        md(block_text)
    
    # Save output
    try:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        print(f"✅ Document saved to: {output_path}")
        return True
    except Exception as e:
        print(f"Error saving document: {e}")
        return False


def main():
    parser = argparse.ArgumentParser(
        description="Convert Markdown with YAML frontmatter to Word document"
    )
    parser.add_argument(
        "--input", "-i",
        required=True,
        help="Input markdown file"
    )
    parser.add_argument(
        "--template", "-t",
        required=True,
        help="Word template file (.docx or .dotx)"
    )
    parser.add_argument(
        "--output", "-o",
        required=True,
        help="Output Word file"
    )
    
    args = parser.parse_args()
    
    input_path = Path(args.input)
    template_path = Path(args.template)
    output_path = Path(args.output)
    
    if not input_path.exists():
        print(f"Error: Input file not found: {input_path}")
        sys.exit(1)
    
    success = convert_markdown_to_docx(input_path, template_path, output_path)
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
